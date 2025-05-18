from docx import Document
from fpdf import FPDF
import datetime

class Proposal:
    def __init__(self, portfolio):
        self.portfolio = portfolio
        self.client_info = {
            'name': '',
            'company': '',
            'email': '',
            'project_title': '',
            'project_description': '',
            'budget': '',
            'timeline': ''
        }
        self.proposal_sections = []

    def set_client_info(self, name, company, email, project_title, project_description, budget, timeline):
        """Set client and project information"""
        self.client_info = {
            'name': name,
            'company': company,
            'email': email,
            'project_title': project_title,
            'project_description': project_description,
            'budget': budget,
            'timeline': timeline
        }

    def add_section(self, title, content):
        """Add a section to the proposal"""
        self.proposal_sections.append({
            'title': title,
            'content': content
        })

    def generate_docx(self, filename):
        """Generate a DOCX proposal"""
        doc = Document()
        
        # Add header
        doc.add_heading('Professional Proposal', 0)
        doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%Y-%m-%d')}")
        
        # Add client information
        doc.add_heading('Client Information', level=1)
        doc.add_paragraph(f"Client Name: {self.client_info['name']}")
        doc.add_paragraph(f"Company: {self.client_info['company']}")
        doc.add_paragraph(f"Project: {self.client_info['project_title']}")
        
        # Add project details
        doc.add_heading('Project Details', level=1)
        doc.add_paragraph(self.client_info['project_description'])
        doc.add_paragraph(f"Budget: {self.client_info['budget']}")
        doc.add_paragraph(f"Timeline: {self.client_info['timeline']}")
        
        # Add custom sections
        for section in self.proposal_sections:
            doc.add_heading(section['title'], level=1)
            doc.add_paragraph(section['content'])
        
        # Add portfolio summary
        doc.add_heading('About Me', level=1)
        portfolio = self.portfolio.get_portfolio_summary()
        doc.add_paragraph(f"Name: {portfolio['personal_info']['name']}")
        doc.add_paragraph(f"Title: {portfolio['personal_info']['title']}")
        
        # Save the document
        doc.save(filename)

    def generate_pdf(self, filename):
        """Generate a PDF proposal"""
        pdf = FPDF()
        pdf.add_page()
        
        # Add header
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, 'Professional Proposal', 0, 1, 'C')
        pdf.set_font('Arial', '', 12)
        pdf.cell(0, 10, f"Date: {datetime.datetime.now().strftime('%Y-%m-%d')}", 0, 1)
        
        # Add client information
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, 'Client Information', 0, 1)
        pdf.set_font('Arial', '', 12)
        pdf.cell(0, 10, f"Client Name: {self.client_info['name']}", 0, 1)
        pdf.cell(0, 10, f"Company: {self.client_info['company']}", 0, 1)
        pdf.cell(0, 10, f"Project: {self.client_info['project_title']}", 0, 1)
        
        # Add project details
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, 'Project Details', 0, 1)
        pdf.set_font('Arial', '', 12)
        pdf.multi_cell(0, 10, self.client_info['project_description'])
        pdf.cell(0, 10, f"Budget: {self.client_info['budget']}", 0, 1)
        pdf.cell(0, 10, f"Timeline: {self.client_info['timeline']}", 0, 1)
        
        # Add custom sections
        for section in self.proposal_sections:
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, section['title'], 0, 1)
            pdf.set_font('Arial', '', 12)
            pdf.multi_cell(0, 10, section['content'])
        
        # Save the PDF
        pdf.output(filename) 